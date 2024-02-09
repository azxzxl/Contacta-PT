#main_logic.py

from docx import Document


def substituir_texto_no_documento_e_preencher_tabelas(caminho_do_template, substituicoes, dados_para_tabela_controle, dados_para_tabela_produtos, dados_para_tabela_servicos):
    doc = Document(caminho_do_template)

    # Substituição de texto nos parágrafos
    for p in doc.paragraphs:
        for key, value in substituicoes.items():
            if key in p.text:
                inline = p.runs
                for i in inline:
                    if key in i.text:
                        i.text = i.text.replace(key, value)

    # Preenchimento da tabela de controle de documento
    tabela_controle = doc.tables[0]  # A primeira tabela é a de controle
    # Inicia no 1 para pular o cabeçalho
    for i, row_data in enumerate(dados_para_tabela_controle, start=1):
        for j, cell_data in enumerate(row_data):
            tabela_controle.cell(i, j).text = str(cell_data)

    # Preenchimento da tabela de produtos
    tabela_produtos = doc.tables[1]  # A segunda tabela é a de produtos
    # Inicia no 1 para pular o cabeçalho
    for i, row_data in enumerate(dados_para_tabela_produtos, start=1):
        for j, cell_data in enumerate(row_data):
            # Insere os dados apenas se a célula estiver vazia
            if tabela_produtos.cell(i, j).text.strip() == '':
                tabela_produtos.cell(i, j).text = str(cell_data)

    # Preenchimento da tabela de serviços
    tabela_servicos = doc.tables[2]  # A terceira tabela é a de serviços
    # Inicia no 1 para pular o cabeçalho
    for i, row_data in enumerate(dados_para_tabela_servicos, start=1):
        for j, cell_data in enumerate(row_data):
            tabela_servicos.cell(i, j).text = str(cell_data)

    doc.save(caminho_do_template.replace('.docx', '_modificado.docx'))


# Caminho do seu documento template
caminho_do_template = r'C:\Users\Lucas Rangel\Desktop\Contacta\Conteudo Contacta\TESTEPT.docx'

# Dicionário com as substituições a serem feitas
substituicoes = {
    "[fabricante]": "FABRICANTE DESEJADO",
    "[solução]": "EX DE SOLUÇAO",
    "[documento]": "DOC AQ",
    "[serviços]": "SERVIÇO AQ",
    "[Produto]": "PRODUTO AQ",
    "[produto]": "PRODUTO AQ",
    "[Texto descritivo]": "Descritivo do produto"
}

# Dados para preencher a tabela de controle
dados_para_tabela_controle = [
    ('1.0', '2024-02-07', 'Primeira versão do documento.'),
    # Adicione mais linhas conforme necessário
]

# Dados para preencher a tabela de produtos nas células vazias
dados_para_tabela_produtos = [
    ('Novo Part Number 1', 'Nova Descrição 1', 'Nova Quantidade 1'),
    ('Novo Part Number 2', 'Nova Descrição 2', 'Nova Quantidade 2'),
    # Adicione mais dados conforme necessário
]

# Dados para preencher a tabela de serviços
dados_para_tabela_servicos = [
    ('1', 'Descrição do serviço 1', 'Part Number 1', '5'),
    # Adicione mais linhas conforme necessário
]

# Chamada da função com os dados fornecidos
substituir_texto_no_documento_e_preencher_tabelas(
    caminho_do_template, substituicoes, dados_para_tabela_controle, dados_para_tabela_produtos, dados_para_tabela_servicos)
